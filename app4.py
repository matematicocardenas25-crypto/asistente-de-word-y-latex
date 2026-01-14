import streamlit as st
from PIL import Image, ImageDraw, ImageOps
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import numpy as np
import io
import json
import re
from datetime import datetime

# --- 1. CONFIGURACIÃ“N DE IDENTIDAD (DEFINICIÃ“N ABSOLUTA AL INICIO) ---
def obtener_fecha_espanol():
Â  Â  meses = {
Â  Â  Â  Â  "January": "Enero", "February": "Febrero", "March": "Marzo", "April": "Abril",
Â  Â  Â  Â  "May": "Mayo", "June": "Junio", "July": "Julio", "August": "Agosto",
Â  Â  Â  Â  "September": "Septiembre", "October": "Octubre", "November": "Noviembre", "December": "Diciembre"
Â  Â  }
Â  Â  ahora = datetime.now()
Â  Â  mes_nombre = meses.get(ahora.strftime('%B'), ahora.strftime('%B'))
Â  Â  return f"{ahora.day} de {mes_nombre}, {ahora.year}"

# Variables de firma globales para evitar NameError en la Vista Previa
fecha_actual = obtener_fecha_espanol()
firma_line1 = "Ismael Antonio Cardenas LÃ³pez"
firma_line2 = "Licenciado en MatemÃ¡tica Unan LeÃ³n Nicaragua"

st.set_page_config(page_title="Sistema Ismael CÃ¡rdenas - UNAN LeÃ³n", layout="wide")

# --- 2. MOTOR DE REDACCIÃ“N ACADÃ‰MICA (LENGUAJE ROBUSTO) ---
def generar_textos_robustos(titulo):
Â  Â  return {
Â  Â  Â  Â  "intro": f"El presente compendio tÃ©cnico, titulado '{titulo}', constituye una sistematizaciÃ³n rigurosa de los fundamentos analÃ­ticos y estructurales de las ciencias exactas. Bajo la autorÃ­a del Lic. Ismael CÃ¡rdenas LÃ³pez, este documento articula la abstracciÃ³n simbÃ³lica con la verificaciÃ³n fenomenolÃ³gica, estableciendo una base sÃ³lida para el pensamiento lÃ³gico-matemÃ¡tico avanzado y garantizando un rigor acadÃ©mico acorde a los mÃ¡s altos estÃ¡ndares institucionales de la UNAN LeÃ³n.",
Â  Â  Â  Â  "conclu": f"Tras el anÃ¡lisis pormenorizado de los elementos expuestos en torno a '{titulo}', se concluye que la convergencia entre el rigor analÃ­tico y la modelizaciÃ³n computacional permite una comprensiÃ³n holÃ­stica de los comportamientos estudiados. La evidencia teÃ³rica aquÃ­ presentada ratifica la importancia de la precisiÃ³n axiomÃ¡tica en la resoluciÃ³n de problemas complejos.",
Â  Â  Â  Â  "recom": "Se recomienda encarecidamente someter los resultados analÃ­ticos a un proceso de contraste crÃ­tico frente a modelos de simulaciÃ³n numÃ©rica para validar su estabilidad. Asimismo, se sugiere profundizar en el estudio de las propiedades intrÃ­nsecas de los marcos teÃ³ricos aquÃ­ abordados, fomentando la aplicaciÃ³n de estos modelos en contextos interdisciplinarios."
Â  Â  }

# --- 3. MOTOR DE LIMPIEZA DE LATEX PARA WORD (ELIMINA $ Y \DOTS) ---
def limpiar_para_word(texto):
Â  Â  if not texto: return ""
Â  Â  # Eliminar sÃ­mbolos de dÃ³lar y puntos suspensivos de LaTeX
Â  Â  texto = texto.replace("$", "").replace(r"\dots", "...").replace(r"\cdots", "...")
Â  Â  # Reemplazos de comandos comunes
Â  Â  reemplazos = {
Â  Â  Â  Â  r"\\left(": "(", r"\\right)": ")", r"\\left[": "[", r"\\right]": "]",
Â  Â  Â  Â  r"\\infty": "infinito", r"\\times": "x", r"\\cdot": "Â·", r"\\": "", r"\,": " "
Â  Â  }
Â  Â  # Traducir fracciones \frac{a}{b} -> (a/b)
Â  Â  texto = re.sub(r'\\frac\{(.*?)\}\{(.*?)\}', r'(\1/\2)', texto)
Â  Â  # Eliminar cualquier comando \comando{...} restante
Â  Â  texto = re.sub(r'\\[a-zA-Z]+\{(.*?)\}', r'\1', texto)
Â  Â  for lat, plain in reemplazos.items():
Â  Â  Â  Â  texto = texto.replace(lat, plain)
Â  Â  return texto.strip()

# --- 4. GESTIÃ“N DE IMAGEN CIRCULAR ---
def preparar_foto_circular():
Â  Â  try:
Â  Â  Â  Â  img = Image.open("foto.png").convert("RGBA")
Â  Â  except:
Â  Â  Â  Â  img = Image.new('RGBA', (400, 400), (255, 255, 255, 0))
Â  Â  Â  Â  draw = ImageDraw.Draw(img)
Â  Â  Â  Â  draw.ellipse((0, 0, 400, 400), fill=(26, 82, 118))
Â  Â Â 
Â  Â  mask = Image.new('L', (400, 400), 0)
Â  Â  ImageDraw.Draw(mask).ellipse((0, 0, 400, 400), fill=255)
Â  Â  output = ImageOps.fit(img, (400, 400), centering=(0.5, 0.5))
Â  Â  output.putalpha(mask)
Â  Â  buf = io.BytesIO()
Â  Â  output.save(buf, format='PNG')
Â  Â  buf.seek(0)
Â  Â  return buf

# --- 5. PERSISTENCIA DE DATOS ---
if 'contenido' not in st.session_state: st.session_state.contenido = ""
if 'ejercicios' not in st.session_state: st.session_state.ejercicios = ""

st.title("ğŸ“ Compilador CientÃ­fico de Ã‰lite - UNAN LeÃ³n")

with st.sidebar:
Â  Â  st.header("ğŸ’¾ Respaldo de Seguridad")
Â  Â  if st.button("ğŸ“¥ Crear Punto de RestauraciÃ³n"):
Â  Â  Â  Â  data_respaldo = {"contenido": st.session_state.contenido, "ejercicios": st.session_state.ejercicios}
Â  Â  Â  Â  st.download_button("Descargar Respaldo (.json)", json.dumps(data_respaldo), "respaldo_ismael.json")

col_in, col_pre = st.columns([1, 1.2])

with col_in:
Â  Â  st.subheader("ğŸ“¥ Panel de Insumos")
Â  Â  titulo_proy = st.text_input("TÃ­tulo del Proyecto", "AnÃ¡lisis y Modelado MatemÃ¡tico")
Â  Â  st.session_state.contenido = st.text_area("Cuerpo del Contenido (LaTeX):", value=st.session_state.contenido, height=350)
Â  Â Â 
Â  Â  st.subheader("ğŸ“Š Motor GrÃ¡fico")
Â  Â  func_in = st.text_input("FunciÃ³n f(x):", "np.sin(x) * np.exp(-x/10)")
Â  Â  buf_graf = io.BytesIO()
Â  Â  try:
Â  Â  Â  Â  x_vals = np.linspace(-10, 20, 1000)
Â  Â  Â  Â  y_vals = eval(func_in, {"x": x_vals, "np": np})
Â  Â  Â  Â  fig, ax = plt.subplots(figsize=(8, 4))
Â  Â  Â  Â  ax.plot(x_vals, y_vals, color='#1A5276', linewidth=2)
Â  Â  Â  Â  ax.grid(True, alpha=0.3)
Â  Â  Â  Â  fig.savefig(buf_graf, format='png', dpi=300); plt.close(fig); buf_graf.seek(0)
Â  Â  except: pass
Â  Â Â 
Â  Â  st.session_state.ejercicios = st.text_area("Ejercicios Propuestos:", value=st.session_state.ejercicios, height=200)

with col_pre:
Â  Â  st.subheader("ğŸ‘ï¸ Vista Previa Institucional")
Â  Â  textos = generar_textos_robustos(titulo_proy)
Â  Â  with st.container(border=True):
Â  Â  Â  Â  st.markdown(f"<div style='text-align: right;'><b>Fecha:</b> {fecha_actual}</div>", unsafe_allow_html=True)
Â  Â  Â  Â  st.markdown(f"<h2 style='text-align:center; color:#1A5276;'>{titulo_proy}</h2>", unsafe_allow_html=True)
Â  Â  Â  Â  st.markdown(f"<p style='text-align:center;'><b>{firma_line1}</b><br><i>{firma_line2}</i></p>", unsafe_allow_html=True)
Â  Â  Â  Â  st.markdown("<hr>", unsafe_allow_html=True)
Â  Â  Â  Â Â 
Â  Â  Â  Â  st.markdown("### 1. IntroducciÃ³n")
Â  Â  Â  Â  st.write(textos['intro'])
Â  Â  Â  Â Â 
Â  Â  Â  Â  st.markdown("### 2. Desarrollo TeÃ³rico")
Â  Â  Â  Â  st.markdown(st.session_state.contenido)
Â  Â  Â  Â Â 
Â  Â  Â  Â  if buf_graf.getbuffer().nbytes > 0:
Â  Â  Â  Â  Â  Â  st.image(buf_graf, caption="AnÃ¡lisis GrÃ¡fico")
Â  Â  Â  Â  Â  Â Â 
Â  Â  Â  Â  st.markdown("### 3. Ejercicios Propuestos")
Â  Â  Â  Â  st.markdown(st.session_state.ejercicios)
Â  Â  Â  Â Â 
Â  Â  Â  Â  st.markdown("### 4. Conclusiones")
Â  Â  Â  Â  st.write(textos['conclu'])
Â  Â  Â  Â Â 
Â  Â  Â  Â  st.markdown("### 5. Recomendaciones")
Â  Â  Â  Â  st.write(textos['recom'])

# --- 6. GENERACIÃ“N DE DOCUMENTOS ---
if st.button("ğŸš€ Compilar DocumentaciÃ³n de Ã‰lite"):
Â  Â  textos = generar_textos_robustos(titulo_proy)
Â  Â  doc = Document()
Â  Â Â 
Â  Â  # Encabezado: Foto Circular a la derecha y Fecha a la izquierda
Â  Â  header_table = doc.add_table(rows=1, cols=2)
Â  Â  header_table.columns[0].width = Inches(4.5)
Â  Â  header_table.cell(0, 0).text = fecha_actual
Â  Â Â 
Â  Â  celda_foto = header_table.cell(0, 1).add_paragraph()
Â  Â  celda_foto.alignment = WD_ALIGN_PARAGRAPH.RIGHT
Â  Â  celda_foto.add_run().add_picture(preparar_foto_circular(), width=Inches(1.0))

Â  Â  # TÃ­tulo y Firma centrada
Â  Â  doc.add_heading('\n' + titulo_proy, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
Â  Â  f1 = doc.add_paragraph(firma_line1)
Â  Â  f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
Â  Â  f2 = doc.add_paragraph(firma_line2)
Â  Â  f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
Â  Â  f2.runs[0].font.italic = True

Â  Â  # Secciones Combinadas con Limpieza Profunda de LaTeX
Â  Â  secciones_word = [
Â  Â  Â  Â  ("I. IntroducciÃ³n", textos['intro']),
Â  Â  Â  Â  ("II. Desarrollo TeÃ³rico", st.session_state.contenido),
Â  Â  Â  Â  ("III. Ejercicios Propuestos", st.session_state.ejercicios),
Â  Â  Â  Â  ("IV. Conclusiones", textos['conclu']),
Â  Â  Â  Â  ("V. Recomendaciones", textos['recom'])
Â  Â  ]

Â  Â  for tit, cont in secciones_word:
Â  Â  Â  Â  doc.add_heading(tit, 1)
Â  Â  Â  Â  texto_limpio = limpiar_para_word(cont)
Â  Â  Â  Â  for linea in texto_limpio.split('\n'):
Â  Â  Â  Â  Â  Â  if linea.strip(): doc.add_paragraph(linea.strip())

Â  Â  if buf_graf.getbuffer().nbytes > 0:
Â  Â  Â  Â  doc.add_picture(buf_graf, width=Inches(5.5))

Â  Â  w_io = io.BytesIO(); doc.save(w_io); w_io.seek(0)
Â  Â Â 
Â  Â  # CÃ³digo LaTeX (Mantiene sus fÃ³rmulas intactas)
Â  Â  latex_code = f"\\documentclass{{article}}\\usepackage[spanish]{{babel}}\\title{{{titulo_proy}}}\\author{{{firma_line1} \\\\ {firma_line2}}}\\begin{{document}}\\maketitle\n\\section{{I. IntroducciÃ³n}}{textos['intro']}\\section{{II. Desarrollo}}{st.session_state.contenido}\\end{{document}}"
Â  Â Â 
Â  Â  st.download_button("â¬‡ï¸ Descargar Word Final", w_io, f"{titulo_proy}.docx")
Â  Â  st.download_button("â¬‡ï¸ Descargar CÃ³digo LaTeX", latex_code, f"{titulo_proy}.tex")
Â  Â  st.success("Â¡DocumentaciÃ³n compilada con Ã©xito!")
